#!/usr/bin/env python3
"""Ingest tài liệu dự án vào knowledge base BM25 offline.

Đầu vào:  data/projects/<project>/...  (PDF/DOCX/XLSX/TXT)
Đầu ra:   data/knowledge_base/<project-slug>/
            - chunks.jsonl   (1 chunk/dòng + metadata)
            - bm25.pkl       (rank_bm25.BM25Okapi đã fit)
            - manifest.json  (thống kê: bao nhiêu file, chunk, file scan cần OCR)

Đặc điểm:
- KHÔNG cần API key — toàn bộ chạy offline.
- Trích text bằng pypdf (nhanh) -> nếu yếu thì fallback pdfplumber.
- Phát hiện PDF scan (rất ít text / trang) -> ghi nhận "needs_ocr" và bỏ qua.
- Chunk theo ký tự (~ 2500 ký tự ≈ 600-800 token) với chồng lấn nhẹ.
- Idempotent: cùng dấu vân tay (size + mtime + path) thì bỏ qua file đã ingest.

Cách chạy:
    cd apps/agent-engine && source .venv/bin/activate
    python ../../scripts/ingest_projects.py \
        --project eurowindow-light-city \
        --groups qa,brochure,phap-ly,bang-gia
"""
from __future__ import annotations

import argparse
import json
import logging
import os
import pickle
import re
import sys
import time
import unicodedata
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Iterable

logging.basicConfig(
    level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s",
)
log = logging.getLogger("ingest")

REPO_ROOT = Path(__file__).resolve().parent.parent
PROJECTS_DIR = REPO_ROOT / "data" / "projects"
KB_DIR = REPO_ROOT / "data" / "knowledge_base"

# Tên thư mục con bên trong "<project>/Eurowindow Light City/" -> nhóm.
# Khớp bằng substring không phân biệt hoa/thường, đã chuẩn hoá dấu.
GROUP_RULES: list[tuple[str, str]] = [
    # (substring trong path đã lowercase + bỏ dấu, group_id)
    ("q&a", "qa"),
    ("qa", "qa"),
    ("brochure", "brochure"),
    ("phap ly", "phap-ly"),
    ("bang gia", "bang-gia"),
    ("gio hang", "bang-gia"),
    ("phieu tinh gia", "bang-gia"),
    ("to roi", "to-roi"),
    ("to gap", "to-gap"),
    ("tai lieu dao tao", "dao-tao"),
    ("nhan dien thuong hieu", "branding"),
    ("anh mat bang", "mat-bang"),
    ("lich truc", "lich-truc"),
    ("link kenh truyen thong", "truyen-thong"),
]

CHUNK_SIZE = 2500   # ký tự — xấp xỉ 600-800 token tiếng Việt
CHUNK_OVERLAP = 250
MIN_TEXT_PER_PAGE = 80  # < ngưỡng này coi như scan thiếu text


# ============================================================
# Helpers
# ============================================================

def strip_accents(s: str) -> str:
    """Bỏ dấu tiếng Việt để khớp substring."""
    nf = unicodedata.normalize("NFD", s)
    return "".join(c for c in nf if unicodedata.category(c) != "Mn")


def classify_group(rel_path: str) -> str:
    key = strip_accents(rel_path.lower())
    for needle, gid in GROUP_RULES:
        if needle in key:
            return gid
    return "khac"


def file_fingerprint(p: Path) -> str:
    st = p.stat()
    return f"{p.name}|{st.st_size}|{int(st.st_mtime)}"


# ============================================================
# Extractors
# ============================================================

@dataclass
class ExtractResult:
    text: str
    pages: int
    needs_ocr: bool
    extractor: str
    error: str = ""


def extract_pdf(path: Path) -> ExtractResult:
    """Thử pypdf trước (nhanh), fallback pdfplumber nếu text yếu."""
    text, pages, err = "", 0, ""
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        pages = len(reader.pages)
        parts = []
        for i, pg in enumerate(reader.pages):
            try:
                t = pg.extract_text() or ""
            except Exception as e:
                t = ""
                err = f"pypdf page {i}: {e}"
            if t.strip():
                parts.append(f"\n\n[Trang {i+1}]\n{t}")
        text = "".join(parts)
    except Exception as e:
        err = f"pypdf: {e}"

    text_per_page = (len(text) / pages) if pages else 0
    if text_per_page < MIN_TEXT_PER_PAGE:
        # Thử pdfplumber như fallback (đôi khi extract tốt hơn)
        try:
            import pdfplumber
            parts = []
            with pdfplumber.open(str(path)) as pdf:
                pages = len(pdf.pages)
                for i, pg in enumerate(pdf.pages):
                    t = pg.extract_text() or ""
                    if t.strip():
                        parts.append(f"\n\n[Trang {i+1}]\n{t}")
            text2 = "".join(parts)
            if len(text2) > len(text):
                text = text2
            text_per_page = (len(text) / pages) if pages else 0
        except Exception as e:
            err = (err + " | " if err else "") + f"pdfplumber: {e}"

    needs_ocr = pages > 0 and text_per_page < MIN_TEXT_PER_PAGE
    return ExtractResult(
        text=text, pages=pages, needs_ocr=needs_ocr,
        extractor="pypdf+pdfplumber", error=err,
    )


def extract_docx(path: Path) -> ExtractResult:
    try:
        import docx
        d = docx.Document(str(path))
        paras = [p.text for p in d.paragraphs if p.text.strip()]
        # Bảng — nối thành dòng
        for tbl in d.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paras.append(" | ".join(cells))
        text = "\n".join(paras)
        return ExtractResult(text=text, pages=1, needs_ocr=False, extractor="python-docx")
    except Exception as e:
        return ExtractResult(text="", pages=0, needs_ocr=False, extractor="python-docx", error=str(e))


def extract_xlsx(path: Path) -> ExtractResult:
    try:
        from openpyxl import load_workbook
        wb = load_workbook(str(path), data_only=True, read_only=True)
        out = []
        for ws in wb.worksheets:
            out.append(f"\n\n[Sheet: {ws.title}]")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    out.append(" | ".join(cells))
        return ExtractResult(text="\n".join(out), pages=1, needs_ocr=False, extractor="openpyxl")
    except Exception as e:
        return ExtractResult(text="", pages=0, needs_ocr=False, extractor="openpyxl", error=str(e))


def extract_txt(path: Path) -> ExtractResult:
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
        return ExtractResult(text=text, pages=1, needs_ocr=False, extractor="plain")
    except Exception as e:
        return ExtractResult(text="", pages=0, needs_ocr=False, extractor="plain", error=str(e))


def extract(path: Path) -> ExtractResult:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_pdf(path)
    if ext == ".docx":
        return extract_docx(path)
    if ext in (".xlsx", ".xls"):
        return extract_xlsx(path)
    if ext == ".txt":
        return extract_txt(path)
    return ExtractResult(text="", pages=0, needs_ocr=False, extractor="none",
                         error=f"unsupported ext {ext}")


# ============================================================
# Chunking
# ============================================================

_WS = re.compile(r"[ \t]+")
_NL = re.compile(r"\n{3,}")

def normalize_text(t: str) -> str:
    t = unicodedata.normalize("NFC", t)
    t = _WS.sub(" ", t)
    t = _NL.sub("\n\n", t)
    return t.strip()


def split_into_chunks(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    text = normalize_text(text)
    if not text:
        return []
    chunks = []
    n = len(text)
    i = 0
    while i < n:
        end = min(i + size, n)
        # Cố gắng cắt ở ranh giới câu/đoạn gần end
        if end < n:
            for delim in ("\n\n", "\n", ". ", "? ", "! "):
                cut = text.rfind(delim, i + size // 2, end)
                if cut != -1:
                    end = cut + len(delim)
                    break
        chunk = text[i:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= n:
            break
        i = max(end - overlap, i + 1)
    return chunks


# ============================================================
# Tokenizer (cho BM25) — chuẩn hoá, lowercase, BỎ DẤU để robust
# ============================================================

_TOKEN = re.compile(r"[a-z0-9]+")

def tokenize(text: str) -> list[str]:
    t = strip_accents(text.lower())
    return _TOKEN.findall(t)


# ============================================================
# Main pipeline
# ============================================================

def find_project_files(project_dir: Path) -> list[Path]:
    exts = {".pdf", ".docx", ".xlsx", ".xls", ".txt"}
    return [p for p in project_dir.rglob("*") if p.is_file() and p.suffix.lower() in exts]


def filter_by_groups(files: list[Path], allowed_groups: set[str] | None, project_dir: Path):
    """Trả list (path, group). Nếu allowed_groups=None thì lấy tất cả."""
    out = []
    for p in files:
        rel = p.relative_to(project_dir).as_posix()
        g = classify_group(rel)
        if allowed_groups is None or g in allowed_groups:
            out.append((p, g, rel))
    return out


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--project", required=True, help="Tên thư mục dự án trong data/projects/")
    ap.add_argument("--groups", default="", help="CSV danh sách group_id; rỗng = tất cả. Ví dụ: qa,brochure,phap-ly,bang-gia")
    ap.add_argument("--rebuild", action="store_true", help="Xoá index cũ và build lại")
    args = ap.parse_args()

    project_dir = PROJECTS_DIR / args.project
    if not project_dir.exists():
        log.error("Không tìm thấy thư mục dự án: %s", project_dir)
        return 1

    allowed = None
    if args.groups.strip():
        allowed = {g.strip() for g in args.groups.split(",") if g.strip()}
        log.info("Lọc theo nhóm: %s", sorted(allowed))
    else:
        log.info("Không lọc nhóm — ingest tất cả file hỗ trợ")

    out_dir = KB_DIR / args.project
    out_dir.mkdir(parents=True, exist_ok=True)

    chunks_path = out_dir / "chunks.jsonl"
    bm25_path = out_dir / "bm25.pkl"
    manifest_path = out_dir / "manifest.json"

    # Idempotent: đọc fingerprint cũ
    existing_fps: set[str] = set()
    existing_chunks: list[dict] = []
    if chunks_path.exists() and not args.rebuild:
        with chunks_path.open("r", encoding="utf-8") as f:
            for line in f:
                d = json.loads(line)
                existing_chunks.append(d)
                existing_fps.add(d["source_fp"])
        log.info("Đã có %d chunk từ trước (%d file). Sẽ chỉ thêm file mới.",
                 len(existing_chunks), len({c["source_fp"] for c in existing_chunks}))

    files = find_project_files(project_dir)
    selected = filter_by_groups(files, allowed, project_dir)
    log.info("Tìm thấy %d file phù hợp", len(selected))

    new_chunks: list[dict] = []
    needs_ocr: list[dict] = []
    extracted_ok: list[dict] = []
    errors: list[dict] = []

    t0 = time.time()
    for path, group, rel in selected:
        fp = file_fingerprint(path)
        if fp in existing_fps and not args.rebuild:
            log.info("[skip] đã ingest: %s", rel)
            continue
        log.info("[%s] %s", group, rel)
        res = extract(path)
        if res.error:
            log.warning("  ! lỗi extract: %s", res.error)
        if res.needs_ocr:
            log.warning("  ! cần OCR (chỉ %d ký tự / %d trang)", len(res.text), res.pages)
            needs_ocr.append({
                "file": rel, "group": group, "pages": res.pages,
                "chars_extracted": len(res.text),
                "reason": "PDF scan — không có text layer",
            })
            continue
        if not res.text.strip():
            log.warning("  ! không có text")
            errors.append({"file": rel, "group": group, "error": res.error or "empty text"})
            continue
        parts = split_into_chunks(res.text)
        log.info("  -> %d ký tự / %d trang / %d chunk", len(res.text), res.pages, len(parts))
        extracted_ok.append({
            "file": rel, "group": group, "pages": res.pages,
            "chars": len(res.text), "chunks": len(parts),
            "extractor": res.extractor,
        })
        for ci, ch in enumerate(parts):
            new_chunks.append({
                "id": f"{fp}::{ci}",
                "project": args.project,
                "group": group,
                "source_file": rel,
                "source_fp": fp,
                "chunk_index": ci,
                "text": ch,
            })

    all_chunks = existing_chunks + new_chunks
    if not all_chunks:
        log.error("Không có chunk nào để build index — dừng")
        return 2

    # Ghi chunks.jsonl
    with chunks_path.open("w", encoding="utf-8") as f:
        for c in all_chunks:
            f.write(json.dumps(c, ensure_ascii=False) + "\n")

    # Build BM25
    from rank_bm25 import BM25Okapi
    log.info("Build BM25 trên %d chunk...", len(all_chunks))
    corpus_tokens = [tokenize(c["text"]) for c in all_chunks]
    bm25 = BM25Okapi(corpus_tokens)
    with bm25_path.open("wb") as f:
        pickle.dump({"bm25": bm25, "tokens_len": [len(t) for t in corpus_tokens]}, f)

    # Manifest
    manifest = {
        "project": args.project,
        "built_at": int(time.time()),
        "total_chunks": len(all_chunks),
        "new_chunks_this_run": len(new_chunks),
        "files_extracted": extracted_ok,
        "files_needs_ocr": needs_ocr,
        "files_error": errors,
        "groups_requested": sorted(allowed) if allowed else None,
        "chunk_size": CHUNK_SIZE,
        "chunk_overlap": CHUNK_OVERLAP,
    }
    with manifest_path.open("w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)

    dt = time.time() - t0
    log.info("Xong sau %.1fs. Index: %s", dt, out_dir)
    log.info("Tổng chunk: %d (mới: %d). File cần OCR: %d. Lỗi: %d.",
             len(all_chunks), len(new_chunks), len(needs_ocr), len(errors))
    return 0


if __name__ == "__main__":
    sys.exit(main())
