"""Trích xuất text + chunk + tokenize cho tài liệu upload (runtime).

Dùng cho Sale Learning Center: khi admin upload PDF/DOCX/XLSX/ảnh, ta trích
text rồi chunk để index BM25. Logic căn chỉnh với scripts/ingest_projects.py
để cùng định dạng chunk/tokenizer (BM25 khớp giữa offline ingest và runtime).

KHÔNG cần API key — chạy offline hoàn toàn. Ảnh (PNG/JPG) chỉ OCR được nếu môi
trường có pytesseract; nếu không, trả text rỗng (tài liệu vẫn lưu + index theo
tiêu đề để tìm theo tên).
"""

from __future__ import annotations

import logging
import re
import unicodedata
from dataclasses import dataclass
from pathlib import Path

log = logging.getLogger(__name__)

CHUNK_SIZE = 2500  # ký tự — xấp xỉ 600-800 token tiếng Việt
CHUNK_OVERLAP = 250
MIN_TEXT_PER_PAGE = 80  # < ngưỡng này coi như PDF scan thiếu text

SUPPORTED_EXTS = {".pdf", ".docx", ".xlsx", ".xls", ".txt", ".png", ".jpg", ".jpeg"}


@dataclass
class ExtractResult:
    text: str
    pages: int
    needs_ocr: bool
    extractor: str
    error: str = ""


# ============================================================
# Chuẩn hoá tiếng Việt
# ============================================================

def strip_accents(s: str) -> str:
    nf = unicodedata.normalize("NFD", s)
    return "".join(c for c in nf if unicodedata.category(c) != "Mn")


_TOKEN = re.compile(r"[a-z0-9]+")


def tokenize(text: str) -> list[str]:
    """Tokenizer cho BM25 — lowercase + bỏ dấu (khớp ingest_projects.py)."""
    return _TOKEN.findall(strip_accents(text.lower()))


_WS = re.compile(r"[ \t]+")
_NL = re.compile(r"\n{3,}")


def normalize_text(t: str) -> str:
    t = unicodedata.normalize("NFC", t)
    t = _WS.sub(" ", t)
    t = _NL.sub("\n\n", t)
    return t.strip()


def split_into_chunks(
    text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP
) -> list[str]:
    """Cắt text thành các đoạn ~size ký tự, ưu tiên ranh giới câu/đoạn."""
    text = normalize_text(text)
    if not text:
        return []
    chunks: list[str] = []
    n = len(text)
    i = 0
    while i < n:
        end = min(i + size, n)
        if end < n:
            for delim in ("\n\n", "\n", ". ", "? ", "! "):
                cut = text.rfind(delim, i + size // 2, end)
                if cut != -1:
                    end = cut + len(delim)
                    break
        chunk = text[i:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= n:
            break
        i = max(end - overlap, i + 1)
    return chunks


# ============================================================
# Extractor theo định dạng
# ============================================================

def _extract_pdf(path: Path) -> ExtractResult:
    text, pages, err = "", 0, ""
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        pages = len(reader.pages)
        parts = []
        for i, pg in enumerate(reader.pages):
            try:
                t = pg.extract_text() or ""
            except Exception as e:  # noqa: BLE001
                t = ""
                err = f"pypdf page {i}: {e}"
            if t.strip():
                parts.append(f"\n\n[Trang {i + 1}]\n{t}")
        text = "".join(parts)
    except Exception as e:  # noqa: BLE001
        err = f"pypdf: {e}"

    text_per_page = (len(text) / pages) if pages else 0
    if text_per_page < MIN_TEXT_PER_PAGE:
        try:
            import pdfplumber

            parts = []
            with pdfplumber.open(str(path)) as pdf:
                pages = len(pdf.pages)
                for i, pg in enumerate(pdf.pages):
                    t = pg.extract_text() or ""
                    if t.strip():
                        parts.append(f"\n\n[Trang {i + 1}]\n{t}")
            text2 = "".join(parts)
            if len(text2) > len(text):
                text = text2
            text_per_page = (len(text) / pages) if pages else 0
        except Exception as e:  # noqa: BLE001
            err = (err + " | " if err else "") + f"pdfplumber: {e}"

    needs_ocr = pages > 0 and text_per_page < MIN_TEXT_PER_PAGE
    return ExtractResult(text, pages, needs_ocr, "pypdf+pdfplumber", err)


def _extract_docx(path: Path) -> ExtractResult:
    try:
        import docx

        d = docx.Document(str(path))
        paras = [p.text for p in d.paragraphs if p.text.strip()]
        for tbl in d.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paras.append(" | ".join(cells))
        return ExtractResult("\n".join(paras), 1, False, "python-docx")
    except Exception as e:  # noqa: BLE001
        return ExtractResult("", 0, False, "python-docx", str(e))


def _extract_xlsx(path: Path) -> ExtractResult:
    try:
        from openpyxl import load_workbook

        wb = load_workbook(str(path), data_only=True, read_only=True)
        out = []
        for ws in wb.worksheets:
            out.append(f"\n\n[Sheet: {ws.title}]")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) for c in row if c is not None]
                if cells:
                    out.append(" | ".join(cells))
        return ExtractResult("\n".join(out), 1, False, "openpyxl")
    except Exception as e:  # noqa: BLE001
        return ExtractResult("", 0, False, "openpyxl", str(e))


def _extract_txt(path: Path) -> ExtractResult:
    try:
        return ExtractResult(
            path.read_text(encoding="utf-8", errors="replace"), 1, False, "plain"
        )
    except Exception as e:  # noqa: BLE001
        return ExtractResult("", 0, False, "plain", str(e))


def _extract_image(path: Path) -> ExtractResult:
    """OCR ảnh nếu có pytesseract; nếu không, trả rỗng (cần OCR)."""
    try:
        import pytesseract  # type: ignore
        from PIL import Image  # type: ignore

        text = pytesseract.image_to_string(Image.open(str(path)), lang="vie+eng")
        return ExtractResult(text or "", 1, not bool(text.strip()), "pytesseract")
    except Exception as e:  # noqa: BLE001
        # Không có OCR engine — không phải lỗi nghiêm trọng, chỉ là cần OCR.
        return ExtractResult("", 1, True, "none", f"ocr unavailable: {e}")


def extract(path: Path) -> ExtractResult:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    if ext in (".xlsx", ".xls"):
        return _extract_xlsx(path)
    if ext == ".txt":
        return _extract_txt(path)
    if ext in (".png", ".jpg", ".jpeg"):
        return _extract_image(path)
    return ExtractResult("", 0, False, "none", f"unsupported ext {ext}")
